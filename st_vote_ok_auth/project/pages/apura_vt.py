# Load important libraries 
import pandas as pd
import streamlit as st
import Lbs001
#######################
# Imports para docx
#######################
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_BREAK
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
###################################

def app():
    st.title("Apuração:")
    pleito = 100
    votos_df = []
    votantes_df = []
    votos = Lbs001.lists_tb('VOTO_RG')
    votantes = Lbs001.lists_tb('VOTANTE_RG')
    # a = []
    # # Carrega os dados da tabela VOTO_RG
    # for linh in votos:
    #     a.append(linh[2])
    # st.text(a)   
    # # Tranforma a lista a em uma lista de itens unicos
    # b = set(a)
    # c = list(b)
    # st.text(c)
    # Separa apenas os votos de um pleito especifico na lista votantes_df
    for eleitores in votantes:
        if eleitores[2] == pleito:
            votantes_df.append(eleitores)
    #dfvt = pd.DataFrame(votantes_df, columns =['Placet', 'Data', 'Pleito'])
    
    # Verifica o placet do votante de votante_df e inclue o nome na lista
    lista_placet_nomes = []
    for lnhs in votantes_df:
        placet = lnhs[0]
        nome = Lbs001.encontra_nome(placet,'USUARIO')
        lnhs1 = lnhs+(nome,)
        lista_placet_nomes.append(lnhs1)
    df3 = pd.DataFrame(lista_placet_nomes, columns =['Placet', 'Data', 'Pleito', 'Nome'])
    
    # Separa apenas os votos de um pleito especifico
    for linhas in votos:
        if linhas[2] == pleito:
            votos_df.append(linhas)
    df = pd.DataFrame(votos_df, columns =['Voto', 'Data', 'Pleito'])
    
    st.subheader('Votantes')
    st.dataframe(df3)
    
    st.subheader('Votos computados')
    st.dataframe(df)
    
    st.subheader('Contagem de votos')
    df_count = df['Voto'].value_counts()
    st.dataframe(df_count)
    
    st.subheader('visualização gráfica de votos')
    st.bar_chart(df_count)
    
    # Gera figura do grafico
    hist = df["Voto"].hist()
    fig = hist.get_figure()
    fig.savefig('grafico.png')
    
    #################################################
    #Contrução do documento de apuraccao.docx
    #################################################
    try:
        Chapa1 = (df_count['Chapa1'])
    except KeyError:
        Chapa1 = '0'
    try:
        Chapa2 = (df_count['Chapa2'])
    except KeyError:
        Chapa2 = '0'
    try:
        Voto_Nulo = (df_count['Voto Nulo'])
    except KeyError:
        Voto_Nulo = '0'
    try:
        Voto_branco = (df_count['Voto em branco'])
    except KeyError:
        Voto_branco = '0'
    
    d = Document('modelo-limpo.docx')
    styles = d.styles
    # Table styles
    #https://www.programmersought.com/article/17032383269/

    count = 0
    #d.add_paragraph('')
    p = d.paragraphs[count]
    p.add_run('Relação de votantes:')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table1 = d.add_table(rows=1, cols=3)
    table1.style='Light List Accent 1'
    #table.allow_autofit = True
    # hdr_ = header, row_ = linha
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Placet'
    hdr_cells[1].text = 'Nome'
    hdr_cells[2].text = 'Data do voto'
    #font.size= Pt(30)
    for linha in lista_placet_nomes:
        row_cells = table1.add_row().cells
        row_cells[0].text = str(linha[0])
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row_cells[1].text = str(linha[3])
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row_cells[2].text = str(linha[1])
        row_cells[2].paragraphs[0].runs[0].font.size = Pt(9)

    count = count + 1
    d.add_paragraph('')
    count = count + 1
    d.add_paragraph('')
    p = d.paragraphs[count]
    p.add_run('Relação de votos:')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2 = d.add_table(rows=1, cols=2)
    table2.style='Light List Accent 1'
    #table.allow_autofit = True
    # hdr_ = header, row_ = linha
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Placet'
    hdr_cells[1].text = 'Data do voto'
    #font.size= Pt(30)
    for linha in votos_df:
        row_cells = table2.add_row().cells
        row_cells[0].text = str(linha[0])
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row_cells[1].text = str(linha[1])
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    count = count + 1
    d.add_paragraph('')
    count = count + 1
    d.add_paragraph('')
    p = d.paragraphs[count]
    p.add_run('Contagem de votos:')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table3 = d.add_table(rows=1, cols=2)
    table3.style='Light List Accent 1'
    #table.allow_autofit = True
    # hdr_ = header, row_ = linha
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'Voto'
    hdr_cells[1].text = 'Qtde.'
    row_cells1 = table3.add_row().cells
    row_cells1[0].text = 'Chapa1'
    row_cells1[0].paragraphs[0].runs[0].font.bold = True
    row_cells1[0].paragraphs[0].runs[0].font.size = Pt(9)
    row_cells1[1].text = str(Chapa1)
    row_cells1[1].paragraphs[0].runs[0].font.size = Pt(9)

    row_cells2 = table3.add_row().cells
    row_cells2[0].text = 'Chapa2'
    row_cells2[0].paragraphs[0].runs[0].font.bold = True
    row_cells2[0].paragraphs[0].runs[0].font.size = Pt(9)
    row_cells2[1].text = str(Chapa2)
    row_cells2[1].paragraphs[0].runs[0].font.size = Pt(9)

    row_cells3 = table3.add_row().cells
    row_cells3[0].text = 'Voto Nulo'
    row_cells3[0].paragraphs[0].runs[0].font.bold = True
    row_cells3[0].paragraphs[0].runs[0].font.size = Pt(9)
    row_cells3[1].text = str(Voto_Nulo)
    row_cells3[1].paragraphs[0].runs[0].font.size = Pt(9)

    row_cells4 = table3.add_row().cells
    row_cells4[0].text = 'Voto em branco'
    row_cells4[0].paragraphs[0].runs[0].font.bold = True
    row_cells4[0].paragraphs[0].runs[0].font.size = Pt(9)
    row_cells4[1].text = str(Voto_branco)
    row_cells4[1].paragraphs[0].runs[0].font.size = Pt(9)
    
    count = count + 1
    d.add_paragraph('')
    count = count + 1
    d.add_paragraph('')
    p = d.paragraphs[count]
    p.add_run('Grafico da contagem de votos:')
    count = count + 1
    d.add_paragraph('')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    d.add_picture('grafico.png', width=Inches(4.0))
    
    d.save('apuracao1.docx')
    
    # Disponibiliza o documento de apuração
    st.text('Faça o download do documento de apuração no link abaixo.')
    st.markdown(Lbs001.get_file('apuracao1.docx', 'Documento de apuração'), unsafe_allow_html=True)  